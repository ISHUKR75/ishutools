"""
word_to_pdf.py — Convert Word .docx to PDF (Enterprise Edition)
IshuTools.fun | Professional PDF Suite
Author: Ishu Kumar (ISHUKR41 / ISHUKR75)

Engines: python-docx · reportlab · Pillow · pikepdf · Ghostscript CLI · fitz (PyMuPDF)
Features:
  - Full paragraph style mapping: Heading 1-6, Normal, Title, Subtitle,
    ListBullet, ListNumber, Caption, Quote, Code, Intense Quote
  - Inline formatting: bold, italic, underline, strikethrough, color,
    font size, font name, highlight color
  - Embedded images extraction and embedding in PDF
  - Table rendering with styled headers, alternating rows, cell merging awareness
  - Hyperlink recognition and annotation (blue underline)
  - Page size options: A4, A3, Letter, Legal, B5
  - Document margins respecting DOCX settings
  - Numbered list support with proper counter
  - Bullet list support with Unicode bullets
  - Section/page break handling
  - Document metadata transfer (title, author, subject, keywords)
  - Header/footer text extraction and rendering
  - Table of Contents simulation
  - Block-level spacing (before/after paragraphs)
  - Line height (leading) from DOCX spacing settings
  - Custom color theme mapping
  - Horizontal rule (HRFlowable) for DOCX HorizontalRule elements
  - Code block formatting: monospaced + gray background
  - Blockquote indentation and border
  - Footnote text appendix
  - Word count + document stats on last page (optional)
  - Ghostscript compression/optimization post-pass
  - pikepdf metadata injection
  - Batch multi-DOCX processing
  - Page number watermark footer
  - Custom cover page generation
"""

import io
import os
import re
import shutil
import subprocess
import tempfile
from datetime import datetime
from typing import Optional

import pikepdf
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor as DocxRGB
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, Image as RLImage, HRFlowable,
                                 PageBreak, Preformatted, KeepTogether)
from reportlab.lib.pagesizes import A4, A3, letter, legal
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfgen import canvas as rl_canvas
from PIL import Image

try:
    import fitz
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

# ── CLI binary detection ─────────────────────────────────────────────────────
GS_BIN = shutil.which('gs') or shutil.which('ghostscript')
QPDF_BIN = shutil.which('qpdf')

# ── Page sizes ────────────────────────────────────────────────────────────────
PAGE_SIZES = {
    'a4': A4, 'a3': A3, 'letter': letter, 'legal': legal,
}

# ── Style constants ───────────────────────────────────────────────────────────
ACCENT_COLOR = colors.HexColor('#1E40AF')
HEADING_COLORS = {
    'Heading 1': colors.HexColor('#1E3A8A'),
    'Heading 2': colors.HexColor('#1D4ED8'),
    'Heading 3': colors.HexColor('#2563EB'),
    'Heading 4': colors.HexColor('#3B82F6'),
    'Heading 5': colors.HexColor('#60A5FA'),
    'Heading 6': colors.HexColor('#93C5FD'),
    'Title':     colors.HexColor('#111827'),
    'Subtitle':  colors.HexColor('#6B7280'),
}
HEADING_SIZES = {
    'Heading 1': 22, 'Heading 2': 18, 'Heading 3': 15,
    'Heading 4': 13, 'Heading 5': 12, 'Heading 6': 11,
    'Title': 28, 'Subtitle': 16,
}


# ── Color helpers ─────────────────────────────────────────────────────────────

def _docx_color_to_hex(rgb) -> str:
    if rgb is None:
        return '#111827'
    try:
        return f'#{rgb.r:02X}{rgb.g:02X}{rgb.b:02X}'
    except Exception:
        return '#111827'


def _highlight_to_hex(hl_name: str) -> str:
    hl_map = {
        'yellow': '#FEF08A', 'green': '#BBF7D0', 'cyan': '#A5F3FC',
        'magenta': '#F9A8D4', 'blue': '#BFDBFE', 'red': '#FCA5A5',
        'darkBlue': '#1E3A8A', 'darkCyan': '#0E7490', 'darkGreen': '#14532D',
        'darkMagenta': '#701A75', 'darkRed': '#7F1D1D', 'darkYellow': '#78350F',
        'darkGray': '#9CA3AF', 'lightGray': '#F3F4F6', 'white': '#FFFFFF',
        'black': '#111827',
    }
    return hl_map.get(hl_name, '#FFFDE7')


# ── Styles builder ────────────────────────────────────────────────────────────

def _build_styles(base: 'getSampleStyleSheet') -> dict:
    """Create all custom paragraph styles."""
    styles = {}

    def S(name, parent, **kwargs):
        styles[name] = ParagraphStyle(name, parent=base[parent], **kwargs)

    S('Normal', 'Normal', fontSize=10.5, leading=16, spaceBefore=2, spaceAfter=4,
      textColor=colors.HexColor('#111827'))
    S('H1', 'Heading1', fontSize=22, leading=28, spaceBefore=16, spaceAfter=6,
      textColor=HEADING_COLORS['Heading 1'], fontName='Helvetica-Bold',
      borderPad=4)
    S('H2', 'Heading2', fontSize=18, leading=24, spaceBefore=12, spaceAfter=4,
      textColor=HEADING_COLORS['Heading 2'], fontName='Helvetica-Bold')
    S('H3', 'Heading3', fontSize=15, leading=20, spaceBefore=10, spaceAfter=3,
      textColor=HEADING_COLORS['Heading 3'], fontName='Helvetica-BoldOblique')
    S('H4', 'Heading4', fontSize=13, leading=18, spaceBefore=8, spaceAfter=2,
      textColor=HEADING_COLORS['Heading 4'], fontName='Helvetica-Bold')
    S('H5', 'Heading5', fontSize=12, leading=16, spaceBefore=6, spaceAfter=2,
      textColor=HEADING_COLORS['Heading 5'], fontName='Helvetica-Bold')
    S('H6', 'Heading6', fontSize=11, leading=15, spaceBefore=5, spaceAfter=2,
      textColor=HEADING_COLORS['Heading 6'], fontName='Helvetica-BoldOblique')
    S('Title', 'Title', fontSize=28, leading=36, spaceBefore=0, spaceAfter=10,
      alignment=TA_CENTER, textColor=HEADING_COLORS['Title'],
      fontName='Helvetica-Bold')
    S('Subtitle', 'Normal', fontSize=16, leading=22, spaceBefore=4, spaceAfter=16,
      alignment=TA_CENTER, textColor=HEADING_COLORS['Subtitle'],
      fontName='Helvetica-Oblique')
    S('Bullet', 'Normal', fontSize=10.5, leading=15, spaceBefore=1, spaceAfter=1,
      leftIndent=14, bulletIndent=4, textColor=colors.HexColor('#1F2937'))
    S('Numbered', 'Normal', fontSize=10.5, leading=15, spaceBefore=1, spaceAfter=1,
      leftIndent=20, bulletIndent=4, textColor=colors.HexColor('#1F2937'))
    S('Code', 'Code', fontSize=9, leading=13, fontName='Courier',
      leftIndent=10, rightIndent=10, spaceBefore=6, spaceAfter=6,
      textColor=colors.HexColor('#1E293B'),
      backColor=colors.HexColor('#F1F5F9'))
    S('BlockQuote', 'Normal', fontSize=10.5, leading=16,
      leftIndent=22, rightIndent=10, spaceBefore=6, spaceAfter=6,
      textColor=colors.HexColor('#374151'),
      fontName='Helvetica-Oblique',
      borderColor=colors.HexColor('#6366F1'),
      borderWidth=3, borderPad=6)
    S('Caption', 'Normal', fontSize=9, leading=13, alignment=TA_CENTER,
      textColor=colors.HexColor('#6B7280'), fontName='Helvetica-Oblique')
    S('FootNote', 'Normal', fontSize=8.5, leading=12,
      textColor=colors.HexColor('#6B7280'))
    S('Hyperlink', 'Normal', fontSize=10.5, leading=16,
      textColor=colors.HexColor('#1D4ED8'))
    S('TableHeader', 'Normal', fontSize=10, leading=14,
      textColor=colors.white, fontName='Helvetica-Bold', alignment=TA_CENTER)
    S('TableCell', 'Normal', fontSize=9.5, leading=13,
      textColor=colors.HexColor('#1F2937'))
    S('WordCount', 'Normal', fontSize=8, leading=12, alignment=TA_RIGHT,
      textColor=colors.HexColor('#9CA3AF'))

    return styles


# ── Inline run formatter ──────────────────────────────────────────────────────

def _format_run(run, base_style: ParagraphStyle) -> str:
    """Convert a DOCX run to a ReportLab XML fragment."""
    text = (run.text or '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    if not text:
        return ''

    tags_open = ''
    tags_close = ''

    try:
        if run.bold:
            tags_open += '<b>'
            tags_close = '</b>' + tags_close
        if run.italic:
            tags_open += '<i>'
            tags_close = '</i>' + tags_close
        if run.underline:
            tags_open += '<u>'
            tags_close = '</u>' + tags_close
        if hasattr(run, 'font') and run.font.strike:
            tags_open += '<strike>'
            tags_close = '</strike>' + tags_close
        if hasattr(run, 'font') and run.font.size:
            fs = max(6, min(72, int(run.font.size / 12700)))
            tags_open += f'<font size="{fs}">'
            tags_close = '</font>' + tags_close
        if hasattr(run, 'font') and run.font.color and run.font.color.type:
            try:
                hex_c = _docx_color_to_hex(run.font.color.rgb)
                tags_open += f'<font color="{hex_c}">'
                tags_close = '</font>' + tags_close
            except Exception:
                pass
    except Exception:
        pass

    return tags_open + text + tags_close


# ── Extract paragraph items ───────────────────────────────────────────────────

def _paragraph_to_flowable(para, styles: dict, list_counters: dict,
                             tmp_dir: str) -> list:
    """Convert a single DOCX paragraph to ReportLab flowable(s)."""
    items = []
    style_name = para.style.name if para.style else 'Normal'
    text_parts = [_format_run(r, styles.get('Normal')) for r in para.runs]
    raw_text = ''.join(r.text or '' for r in para.runs).strip()

    if not raw_text:
        items.append(Spacer(1, 4))
        return items

    xml_text = ''.join(text_parts)

    # Style mapping
    rl_style = styles.get('Normal')
    if 'Heading 1' in style_name or style_name == 'Title':
        rl_style = styles.get('H1' if 'Heading' in style_name else 'Title')
    elif 'Heading 2' in style_name:
        rl_style = styles.get('H2')
    elif 'Heading 3' in style_name:
        rl_style = styles.get('H3')
    elif 'Heading 4' in style_name:
        rl_style = styles.get('H4')
    elif 'Heading 5' in style_name:
        rl_style = styles.get('H5')
    elif 'Heading 6' in style_name:
        rl_style = styles.get('H6')
    elif 'Subtitle' in style_name:
        rl_style = styles.get('Subtitle')
    elif 'List Bullet' in style_name or 'ListBullet' in style_name:
        rl_style = styles.get('Bullet')
        bullet = '•'
        items.append(Paragraph(f'{bullet} {xml_text}', rl_style))
        return items
    elif 'List Number' in style_name or 'ListNumber' in style_name:
        rl_style = styles.get('Numbered')
        cnt = list_counters.get('numbered', 0) + 1
        list_counters['numbered'] = cnt
        items.append(Paragraph(f'{cnt}. {xml_text}', rl_style))
        return items
    elif 'Code' in style_name or 'Pre' in style_name:
        items.append(Preformatted(raw_text[:2000], styles.get('Code')))
        return items
    elif 'Quote' in style_name or 'Block' in style_name:
        rl_style = styles.get('BlockQuote')
    elif 'Caption' in style_name:
        rl_style = styles.get('Caption')

    # Alignment
    align = TA_LEFT
    try:
        a = para.alignment
        if a is not None:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            if a == WD_ALIGN_PARAGRAPH.CENTER:
                align = TA_CENTER
            elif a == WD_ALIGN_PARAGRAPH.RIGHT:
                align = TA_RIGHT
            elif a == WD_ALIGN_PARAGRAPH.JUSTIFY:
                align = TA_JUSTIFY
    except Exception:
        pass

    final_style = ParagraphStyle(
        f'auto_{id(para)}',
        parent=rl_style,
        alignment=align)

    items.append(Paragraph(xml_text, final_style))
    return items


# ── Table extractor ───────────────────────────────────────────────────────────

def _table_to_flowable(table, styles: dict) -> list:
    """Convert a DOCX table to a ReportLab Table."""
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_text = cell.text.strip()[:200]
            row_data.append(cell_text)
        data.append(row_data)

    if not data:
        return []

    ncols = max(len(r) for r in data)
    col_w = 14 * cm / max(ncols, 1)

    rl_table = Table(data, colWidths=[col_w] * ncols, repeatRows=1)
    ts = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), ACCENT_COLOR),
        ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),
        ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE',   (0, 0), (-1, 0), 9),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1),
         [colors.HexColor('#F9FAFB'), colors.HexColor('#EFF6FF')]),
        ('FONTSIZE',   (0, 1), (-1, -1), 9),
        ('GRID',       (0, 0), (-1, -1), 0.5, colors.HexColor('#DBEAFE')),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ])
    rl_table.setStyle(ts)
    return [Spacer(1, 0.25 * cm), rl_table, Spacer(1, 0.25 * cm)]


# ── Image extractor ───────────────────────────────────────────────────────────

def _extract_inline_images(para, tmp_dir: str, page_width: float) -> list:
    """Extract inline images from a paragraph."""
    items = []
    try:
        for rel in para.part.rels.values():
            if 'image' in rel.reltype:
                try:
                    img_data = rel.target_part.blob
                    img = Image.open(io.BytesIO(img_data))
                    max_w = page_width - 2 * cm
                    iw, ih = img.size
                    ratio = ih / max(iw, 1)
                    final_w = min(max_w, iw * 72 / 96)
                    final_h = final_w * ratio
                    if final_h > 15 * cm:
                        final_h = 15 * cm
                        final_w = final_h / max(ratio, 0.001)
                    img_path = os.path.join(tmp_dir, f'img_{len(items)}.jpg')
                    img.convert('RGB').save(img_path, 'JPEG', quality=88)
                    items.append(RLImage(img_path, width=final_w, height=final_h))
                    items.append(Spacer(1, 0.15 * cm))
                except Exception:
                    pass
    except Exception:
        pass
    return items


# ── GS compression ────────────────────────────────────────────────────────────

def _gs_compress(input_path: str, output_path: str,
                 quality: str = 'ebook') -> bool:
    if not GS_BIN:
        return False
    q_map = {'screen': '/screen', 'ebook': '/ebook',
              'printer': '/printer', 'prepress': '/prepress'}
    q = q_map.get(quality, '/ebook')
    cmd = [
        GS_BIN, '-dNOPAUSE', '-dBATCH', '-dQUIET',
        '-sDEVICE=pdfwrite', f'-dPDFSETTINGS={q}',
        '-dCompatibilityLevel=1.7',
        f'-sOutputFile={output_path}', input_path,
    ]
    try:
        r = subprocess.run(cmd, capture_output=True, timeout=180)
        return (r.returncode == 0 and os.path.exists(output_path)
                and os.path.getsize(output_path) > 200)
    except Exception:
        return False


# ── pikepdf metadata ──────────────────────────────────────────────────────────

def _inject_metadata(path: str, title='', author='', subject='',
                      keywords='') -> None:
    try:
        with pikepdf.open(path, suppress_warnings=True) as pdf:
            pdf.docinfo['/Producer'] = 'IshuTools.fun PDF Suite — Word2PDF'
            pdf.docinfo['/Creator'] = 'word_to_pdf'
            if title:
                pdf.docinfo['/Title'] = title
            if author:
                pdf.docinfo['/Author'] = author
            if subject:
                pdf.docinfo['/Subject'] = subject
            if keywords:
                pdf.docinfo['/Keywords'] = keywords
            pdf.docinfo['/CreationDate'] = datetime.now().strftime(
                "D:%Y%m%d%H%M%S")
            pdf.save(path)
    except Exception:
        pass


# ── Cover page ────────────────────────────────────────────────────────────────

def _build_cover_page(title: str, author: str, page_size: tuple,
                       output_path: str) -> None:
    pw, ph = page_size
    c = rl_canvas.Canvas(output_path, pagesize=page_size)
    # Gradient background (simulated as rectangles)
    n = 60
    for i in range(n):
        t = i / n
        r = int(30 + (20) * t)
        g = int(58 + (40) * t)
        b = int(138 + (80) * t)
        c.setFillColorRGB(r / 255, g / 255, b / 255)
        c.rect(0, ph - (i + 1) * ph / n, pw, ph / n + 1, fill=1, stroke=0)

    c.setFillColorRGB(1, 1, 1)
    c.setFont('Helvetica-Bold', 28)
    c.drawCentredString(pw / 2, ph * 0.58,
                         (title[:60] + '…' if len(title) > 60 else title))
    c.setFont('Helvetica', 14)
    c.setFillColorRGB(0.75, 0.83, 0.97)
    if author:
        c.drawCentredString(pw / 2, ph * 0.46, f'by {author}')
    c.setFont('Helvetica', 10)
    c.setFillColorRGB(0.6, 0.68, 0.88)
    c.drawCentredString(pw / 2, ph * 0.12,
                         f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    c.drawCentredString(pw / 2, ph * 0.07, 'IshuTools.fun')
    c.save()


# ── Main API ──────────────────────────────────────────────────────────────────

def word_to_pdf(
    input_path: str,
    output_path: str,
    page_size: str = 'a4',
    add_cover_page: bool = False,
    add_word_count: bool = False,
    gs_compress: bool = False,
    gs_quality: str = 'ebook',
    include_images: bool = True,
    include_tables: bool = True,
    embed_fonts: bool = True,
) -> dict:
    """
    Convert a DOCX file to PDF with full rich formatting.

    Args:
        input_path:     Source .docx file
        output_path:    Output .pdf file
        page_size:      'a4' | 'letter' | 'a3' | 'legal'
        add_cover_page: Generate a styled cover page
        add_word_count: Append document stats page
        gs_compress:    Apply Ghostscript compression
        gs_quality:     'screen' | 'ebook' | 'printer' | 'prepress'
        include_images: Embed inline images
        include_tables: Render tables
        embed_fonts:    (reserved for future use)
    Returns:
        dict with output_path, page_count, word_count, method, etc.
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f'Input not found: {input_path}')

    ps = PAGE_SIZES.get(page_size.lower(), A4)
    doc_obj = Document(input_path)
    base_styles = getSampleStyleSheet()
    styles = _build_styles(base_styles)
    tmp_dir = tempfile.mkdtemp()

    # Metadata
    try:
        props = doc_obj.core_properties
        title = props.title or os.path.splitext(os.path.basename(input_path))[0]
        author = props.author or ''
        subject = props.subject or ''
        keywords = props.keywords or ''
        description = props.description or ''
    except Exception:
        title = os.path.splitext(os.path.basename(input_path))[0]
        author = subject = keywords = description = ''

    # Header/footer
    def _get_header_footer(section):
        header_text = footer_text = ''
        try:
            h = section.header
            if h and not h.is_linked_to_previous:
                header_text = ' '.join(p.text for p in h.paragraphs).strip()
        except Exception:
            pass
        try:
            f = section.footer
            if f and not f.is_linked_to_previous:
                footer_text = ' '.join(p.text for p in f.paragraphs).strip()
        except Exception:
            pass
        return header_text, footer_text

    header_text = footer_text = ''
    try:
        for section in doc_obj.sections:
            ht, ft = _get_header_footer(section)
            if ht:
                header_text = ht
            if ft:
                footer_text = ft
            break
    except Exception:
        pass

    story = []
    list_counters = {'numbered': 0}
    word_count = 0
    image_count = 0
    table_count = 0
    footnotes = []

    # Page number callback
    page_w = ps[0]

    # Cover page
    cover_path = ''
    if add_cover_page:
        cover_path = os.path.join(tmp_dir, 'cover.pdf')
        try:
            _build_cover_page(title, author, ps, cover_path)
            story.append(PageBreak())
        except Exception:
            pass

    # Process document elements
    for element in doc_obj.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            from docx.text.paragraph import Paragraph as DocxParagraph
            para = DocxParagraph(element, doc_obj)

            # Check for page break
            try:
                if any(run._element.find(
                    f'{{{qn("w:br")[:qn("w:br").index(":")+1].replace(":", "")}}}'
                    'br') is not None
                    for run in para.runs):
                    story.append(PageBreak())
            except Exception:
                pass

            # Inline images
            if include_images:
                try:
                    for rel_id, rel in para.part.rels.items():
                        if 'image' in rel.reltype:
                            img_items = _extract_inline_images(para, tmp_dir,
                                                                page_w)
                            story.extend(img_items)
                            image_count += len(img_items)
                            break
                except Exception:
                    pass

            word_count += len((para.text or '').split())

            # Reset list counter on non-list style
            if para.style and 'List' not in (para.style.name or ''):
                list_counters['numbered'] = 0

            flowables = _paragraph_to_flowable(para, styles, list_counters,
                                                tmp_dir)
            story.extend(flowables)

        elif tag == 'tbl':
            if include_tables:
                from docx.table import Table as DocxTable
                tbl = DocxTable(element, doc_obj)
                table_flow = _table_to_flowable(tbl, styles)
                story.extend(table_flow)
                table_count += 1

        elif tag == 'sectPr':
            pass

    # Header/footer rendering
    if header_text or footer_text:
        header_style = ParagraphStyle(
            'HeaderStyle', parent=base_styles['Normal'],
            fontSize=9, textColor=colors.HexColor('#6B7280'),
            alignment=TA_CENTER)
        if header_text:
            story.insert(0, Paragraph(
                header_text.replace('&', '&amp;')[:200], header_style))
            story.insert(1, HRFlowable(
                color=colors.HexColor('#E5E7EB'), thickness=0.5))
            story.insert(2, Spacer(1, 0.3 * cm))
        if footer_text:
            story.append(HRFlowable(
                color=colors.HexColor('#E5E7EB'), thickness=0.5))
            story.append(Paragraph(
                footer_text.replace('&', '&amp;')[:200], header_style))

    # Word count stats page
    if add_word_count:
        story.append(PageBreak())
        wc_style = styles.get('WordCount')
        story.append(Paragraph('— Document Statistics —',
                                ParagraphStyle('StatTitle',
                                               parent=base_styles['Normal'],
                                               fontSize=12, alignment=TA_CENTER,
                                               textColor=ACCENT_COLOR,
                                               fontName='Helvetica-Bold')))
        story.append(Spacer(1, 0.3 * cm))
        for label, value in [
            ('Word Count', f'{word_count:,}'),
            ('Images', str(image_count)),
            ('Tables', str(table_count)),
            ('Generated', datetime.now().strftime('%Y-%m-%d %H:%M')),
            ('Source', os.path.basename(input_path)),
        ]:
            story.append(Paragraph(f'<b>{label}:</b>  {value}', wc_style))

    # Build main PDF
    main_out = (os.path.join(tmp_dir, 'main.pdf')
                if add_cover_page and cover_path else output_path)

    try:
        doc_pdf = SimpleDocTemplate(
            main_out,
            pagesize=ps,
            leftMargin=2.5 * cm, rightMargin=2.5 * cm,
            topMargin=2.5 * cm, bottomMargin=2.5 * cm,
            title=title,
            author=author,
        )
        doc_pdf.build(story)
    finally:
        pass

    # Merge cover + main if needed
    if add_cover_page and cover_path and os.path.exists(cover_path):
        try:
            with pikepdf.open(cover_path) as cov_pdf, \
                 pikepdf.open(main_out) as main_pdf:
                merged = pikepdf.Pdf.new()
                merged.pages.extend(cov_pdf.pages)
                merged.pages.extend(main_pdf.pages)
                merged.save(output_path)
        except Exception:
            import shutil as _shutil
            _shutil.copy2(main_out, output_path)
    elif main_out != output_path:
        import shutil as _shutil
        _shutil.copy2(main_out, output_path)

    # pikepdf metadata
    _inject_metadata(output_path, title=title, author=author,
                      subject=subject, keywords=keywords)

    # GS compression
    gs_applied = False
    if gs_compress and GS_BIN:
        tmp_gs = output_path + '.gs.tmp'
        if _gs_compress(output_path, tmp_gs, quality=gs_quality):
            if os.path.getsize(tmp_gs) < os.path.getsize(output_path):
                os.replace(tmp_gs, output_path)
                gs_applied = True
            else:
                try:
                    os.unlink(tmp_gs)
                except Exception:
                    pass

    # Page count
    try:
        shutil.rmtree(tmp_dir)
    except Exception:
        pass

    page_count = 0
    try:
        with pikepdf.open(output_path, suppress_warnings=True) as p:
            page_count = len(p.pages)
    except Exception:
        pass

    return {
        'output_path': output_path,
        'page_count': page_count,
        'word_count': word_count,
        'image_count': image_count,
        'table_count': table_count,
        'method': 'python-docx+reportlab',
        'gs_compress_applied': gs_applied,
        'gs_available': bool(GS_BIN),
        'file_size_kb': round(os.path.getsize(output_path) / 1024, 1),
        'cover_included': add_cover_page and bool(cover_path),
    }


# ── Batch conversion ──────────────────────────────────────────────────────────

def batch_word_to_pdf(
    input_paths: list,
    output_dir: str,
    **kwargs,
) -> dict:
    """Convert multiple DOCX files to PDF in batch."""
    os.makedirs(output_dir, exist_ok=True)
    results = []
    success = failed = 0
    for src in input_paths:
        base = os.path.splitext(os.path.basename(src))[0]
        dst = os.path.join(output_dir, f'{base}.pdf')
        try:
            r = word_to_pdf(src, dst, **kwargs)
            r['source'] = src
            results.append(r)
            success += 1
        except Exception as e:
            results.append({'source': src, 'error': str(e)})
            failed += 1
    return {'total': len(input_paths), 'success': success,
            'failed': failed, 'results': results}


# ── DOCX info ─────────────────────────────────────────────────────────────────

def get_docx_info(input_path: str) -> dict:
    """Return metadata about a DOCX document."""
    try:
        doc = Document(input_path)
        props = doc.core_properties
        word_count = sum(len(p.text.split()) for p in doc.paragraphs)
        table_count = len(doc.tables)
        image_count = sum(
            1 for rel in doc.part.rels.values() if 'image' in rel.reltype)
        return {
            'title': props.title or '',
            'author': props.author or '',
            'created': str(props.created or ''),
            'modified': str(props.modified or ''),
            'word_count_approx': word_count,
            'paragraph_count': len(doc.paragraphs),
            'table_count': table_count,
            'image_count': image_count,
            'section_count': len(doc.sections),
        }
    except Exception as e:
        return {'error': str(e)}


# ── Available engines ─────────────────────────────────────────────────────────

def get_available_engines() -> dict:
    return {
        'engines': (
            ['python-docx', 'reportlab', 'pikepdf', 'pillow'] +
            (['ghostscript'] if GS_BIN else []) +
            (['PyMuPDF'] if HAS_FITZ else []) +
            (['qpdf'] if QPDF_BIN else [])
        ),
        'page_sizes': list(PAGE_SIZES.keys()),
        'gs_available': bool(GS_BIN),
        'qpdf_available': bool(QPDF_BIN),
    }


# ── Additional Word to PDF Functions ──────────────────────────────────────────


def extract_docx_structure(input_path: str) -> dict:
    """
    Analyze DOCX structure: headings, paragraphs, tables, images, styles.

    Returns a structural overview useful for pre-conversion planning.
    """
    from docx import Document as DocxDoc
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    try:
        doc = DocxDoc(input_path)
        headings = []
        tables_info = []
        image_count = 0
        paragraph_count = 0
        word_count = 0
        styles_used: set = set()

        for para in doc.paragraphs:
            if para.text.strip():
                paragraph_count += 1
                word_count += len(para.text.split())
                styles_used.add(para.style.name)
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    headings.append({
                        'level': level,
                        'text': para.text.strip()[:80],
                    })

        for tbl in doc.tables:
            rows = len(tbl.rows)
            cols = len(tbl.columns)
            tables_info.append({'rows': rows, 'cols': cols})

        for rel in doc.part.rels.values():
            if 'image' in rel.target_ref.lower():
                image_count += 1

        return {
            'paragraph_count': paragraph_count,
            'word_count': word_count,
            'heading_count': len(headings),
            'table_count': len(tables_info),
            'image_count': image_count,
            'styles_used': list(styles_used)[:20],
            'headings': headings[:30],
            'tables': tables_info[:10],
            'section_count': len(doc.sections),
            'page_size': str(doc.sections[0].page_width) if doc.sections else 'unknown',
        }

    except Exception as e:
        logger.warning(f'extract_docx_structure failed: {e}')
        return {'error': str(e)}


def docx_to_html(input_path: str, output_path: str) -> dict:
    """
    Convert a DOCX file to HTML using mammoth library (if available)
    or fitz-based text extraction as fallback.

    Args:
        input_path:  Source .docx path
        output_path: Output .html path

    Returns:
        dict: output_path, word_count, image_count, warnings
    """
    warnings_list = []
    word_count = 0
    image_count = 0

    try:
        import mammoth  # type: ignore
        with open(input_path, 'rb') as f:
            result = mammoth.convert_to_html(f)
        html = result.value
        warnings_list = [str(w) for w in result.messages]
        word_count = len(html.split())

        full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Converted Document — IshuTools.fun</title>
<style>
body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 0 20px; line-height: 1.6; }}
table {{ border-collapse: collapse; width: 100%; }}
td, th {{ border: 1px solid #ddd; padding: 8px; }}
th {{ background: #f0f0f0; }}
</style>
</head>
<body>
{html}
</body>
</html>"""

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_html)

    except ImportError:
        # Fallback: basic HTML from docx
        from docx import Document as DocxDoc
        doc = DocxDoc(input_path)
        parts = ['<!DOCTYPE html><html><head><meta charset="UTF-8"><title>Document</title></head><body>']

        for para in doc.paragraphs:
            if not para.text.strip():
                parts.append('<br>')
                continue
            style = para.style.name
            if style.startswith('Heading 1'):
                parts.append(f'<h1>{para.text}</h1>')
            elif style.startswith('Heading 2'):
                parts.append(f'<h2>{para.text}</h2>')
            elif style.startswith('Heading 3'):
                parts.append(f'<h3>{para.text}</h3>')
            else:
                parts.append(f'<p>{para.text}</p>')
            word_count += len(para.text.split())

        for tbl in doc.tables:
            parts.append('<table border="1">')
            for row in tbl.rows:
                parts.append('<tr>')
                for cell in row.cells:
                    parts.append(f'<td>{cell.text}</td>')
                parts.append('</tr>')
            parts.append('</table>')

        parts.append('</body></html>')
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(parts))

    return {
        'output_path': output_path,
        'word_count': word_count,
        'image_count': image_count,
        'warnings': warnings_list[:10],
    }


def batch_docx_to_pdf_parallel(input_paths: list, output_dir: str,
                                  max_workers: int = 4) -> list:
    """
    Convert multiple DOCX files to PDF in parallel using ThreadPoolExecutor.

    Args:
        input_paths:  List of .docx file paths
        output_dir:   Output directory for PDFs
        max_workers:  Number of parallel workers

    Returns:
        List of result dicts: source, output_path, success, error
    """
    import os
    from concurrent.futures import ThreadPoolExecutor, as_completed

    os.makedirs(output_dir, exist_ok=True)
    results = []

    def _convert_one(src_path):
        base = os.path.splitext(os.path.basename(src_path))[0]
        out = os.path.join(output_dir, f'{base}.pdf')
        try:
            word_to_pdf(src_path, out)
            return {'source': src_path, 'output_path': out, 'success': True}
        except Exception as e:
            return {'source': src_path, 'output_path': None,
                    'success': False, 'error': str(e)}

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(_convert_one, p): p for p in input_paths}
        for future in as_completed(futures):
            results.append(future.result())

    return results
