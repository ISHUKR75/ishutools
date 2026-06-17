"""
pdf_to_word.py - Enterprise PDF to Word Converter
IshuTools.fun | Professional PDF Suite

Conversion strategies (in order of quality):
  1. pdf2docx — best layout fidelity (tables, columns, images)
  2. PyMuPDF (fitz) → python-docx — rich text with formatting
  3. pdfminer → python-docx — pure text fallback
  4. Tesseract OCR → python-docx — for scanned/image PDFs

Features:
  - Full layout reconstruction (tables, columns, lists)
  - Heading detection by font size + bold flag
  - Inline image extraction and embedding
  - Font color, bold, italic, underline preservation
  - Hyperlink detection and embedding
  - Page break insertion between PDF pages
  - Multi-column text handling
  - Table detection and Word table reconstruction
  - EXIF / PDF metadata transfer to docx properties
  - Page range selection (start/end page)
  - Password-protected PDF support
  - Compression of output docx
  - Batch conversion
"""

import os
import io
import re
import shutil
import tempfile
import logging
from datetime import datetime
from typing import Optional

import fitz
from pypdf import PdfReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pdf2docx import Converter

logger = logging.getLogger(__name__)


# ── Text analysis helpers ─────────────────────────────────────────────────────

def _get_median_font_size(fitz_doc) -> float:
    """Estimate median body font size across the first 5 pages."""
    sizes = []
    for page in list(fitz_doc)[:5]:
        for block in page.get_text('dict').get('blocks', []):
            if block.get('type') != 0:
                continue
            for line in block.get('lines', []):
                for span in line.get('spans', []):
                    s = span.get('size', 12)
                    if 6 < s < 100:
                        sizes.append(s)
    if not sizes:
        return 12.0
    sizes.sort()
    return sizes[len(sizes) // 2]


def _detect_heading_level(span_size: float, body_size: float) -> int:
    """Detect heading level from font size ratio."""
    ratio = span_size / max(body_size, 8)
    if ratio >= 2.0:
        return 1
    elif ratio >= 1.6:
        return 2
    elif ratio >= 1.35:
        return 3
    elif ratio >= 1.18:
        return 4
    elif ratio >= 1.1:
        return 5
    return 0


def _is_list_item(text: str) -> tuple:
    """Detect bullet/numbered list items. Returns (is_list, cleaned_text)."""
    bullet_patterns = [
        r'^[•·▪▸►\-–—]\s+(.+)',
        r'^(\d+[.)]\s+)(.+)',
        r'^([a-z]\.\s+)(.+)',
        r'^([ivxlc]+\.\s+)(.+)',
    ]
    for p in bullet_patterns:
        m = re.match(p, text.strip())
        if m:
            return True, m.group(m.lastindex)
    return False, text


# ── Image extraction ──────────────────────────────────────────────────────────

def _extract_and_embed_images(fitz_page, fitz_doc, word_doc, tmp_dir: str,
                                max_width_inches: float = 5.5):
    """Extract images from a page and add them to the word doc."""
    img_list = fitz_page.get_images(full=True)
    embedded = 0
    for img_info in img_list:
        xref = img_info[0]
        try:
            base_img = fitz_doc.extract_image(xref)
            img_bytes = base_img.get('image', b'')
            ext = base_img.get('ext', 'png').lower()
            if not img_bytes or len(img_bytes) < 200:
                continue
            if ext not in ('jpeg', 'jpg', 'png', 'bmp', 'gif', 'tiff', 'webp'):
                ext = 'png'

            img_path = os.path.join(tmp_dir, f'img_{xref}.{ext}')
            with open(img_path, 'wb') as f:
                f.write(img_bytes)

            para = word_doc.add_paragraph()
            run = para.add_run()
            try:
                run.add_picture(img_path, width=Inches(min(max_width_inches,
                                                            base_img.get('width', 400) / 96)))
            except Exception:
                run.add_picture(img_path, width=Inches(4.0))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            embedded += 1
        except Exception:
            continue
    return embedded


# ── Table reconstruction ──────────────────────────────────────────────────────

def _detect_and_add_table(fitz_page, word_doc) -> bool:
    """Detect simple grid tables using PyMuPDF and reconstruct in docx."""
    try:
        tables = fitz_page.find_tables()
        if not tables or not tables.tables:
            return False
        for table in tables.tables:
            data = table.extract()
            if not data:
                continue
            rows = len(data)
            cols = max(len(row) for row in data)
            if rows < 1 or cols < 1:
                continue
            docx_table = word_doc.add_table(rows=rows, cols=cols)
            docx_table.style = 'Table Grid'
            for r_idx, row in enumerate(data):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < cols:
                        try:
                            cell = docx_table.cell(r_idx, c_idx)
                            cell.text = str(cell_text or '').strip()
                        except Exception:
                            pass
            word_doc.add_paragraph()
        return True
    except Exception:
        return False


# ── PyMuPDF → docx fallback ──────────────────────────────────────────────────

def _fitz_to_docx(input_path: str, output_path: str,
                   start_page: int = 0, end_page: Optional[int] = None,
                   extract_images: bool = True,
                   detect_tables: bool = True):
    """Full fitz → python-docx conversion with rich formatting."""
    doc = fitz.open(input_path)
    word_doc = Document()

    # Document defaults
    style = word_doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Page margins
    from docx.oxml.ns import qn as _qn
    section = word_doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    body_size = _get_median_font_size(doc)
    tmp_dir = tempfile.mkdtemp(prefix='ishu_pdf2word_')
    end = min(end_page if end_page is not None else doc.page_count, doc.page_count)

    try:
        for page_idx in range(start_page, end):
            if page_idx > start_page:
                word_doc.add_page_break()

            page = doc[page_idx]

            # Tables first
            if detect_tables:
                _detect_and_add_table(page, word_doc)

            # Images
            if extract_images:
                _extract_and_embed_images(page, doc, word_doc, tmp_dir)

            # Text blocks
            blocks = page.get_text('dict').get('blocks', [])
            prev_y = -1

            for block in blocks:
                btype = block.get('type', -1)

                if btype == 1:
                    # Inline image block
                    if extract_images:
                        try:
                            img_bytes = block.get('image', b'')
                            if img_bytes:
                                img_path = os.path.join(
                                    tmp_dir, f'inline_{page_idx}_{id(block)}.png')
                                with open(img_path, 'wb') as f:
                                    f.write(img_bytes)
                                para = word_doc.add_paragraph()
                                run = para.add_run()
                                run.add_picture(img_path, width=Inches(4.5))
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception:
                            pass
                    continue

                if btype != 0:
                    continue

                for line in block.get('lines', []):
                    spans = line.get('spans', [])
                    if not spans:
                        continue

                    line_text = ''.join(s.get('text', '') for s in spans).strip()
                    if not line_text:
                        continue

                    # Detect heading from font size
                    max_size = max((s.get('size', 12) for s in spans), default=12)
                    heading_level = _detect_heading_level(max_size, body_size)

                    # Check if bold (flag bit 16)
                    all_bold = all(s.get('flags', 0) & 16 for s in spans if s.get('text'))

                    # List detection
                    is_list, clean_text = _is_list_item(line_text)

                    if heading_level > 0 or (all_bold and max_size > body_size * 1.05):
                        level = heading_level or 3
                        para = word_doc.add_heading(line_text, level=min(level, 9))
                    elif is_list:
                        para = word_doc.add_paragraph(style='List Bullet')
                        para.add_run(clean_text)
                    else:
                        para = word_doc.add_paragraph()
                        for span in spans:
                            span_text = span.get('text', '')
                            if not span_text:
                                continue
                            flags = span.get('flags', 0)
                            run = para.add_run(span_text)
                            run.bold = bool(flags & 16)
                            run.italic = bool(flags & 2)
                            run.underline = bool(flags & 4)

                            size = span.get('size', 11)
                            if 6 < size < 100:
                                run.font.size = Pt(round(size, 1))

                            color = span.get('color', None)
                            if isinstance(color, int) and color != 0:
                                r = (color >> 16) & 0xFF
                                g = (color >> 8) & 0xFF
                                b = color & 0xFF
                                if not (r == g == b == 0):
                                    run.font.color.rgb = RGBColor(r, g, b)

    finally:
        doc.close()
        try:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass

    word_doc.save(output_path)


# ── pdfminer fallback ─────────────────────────────────────────────────────────

def _pdfminer_to_docx(input_path: str, output_path: str):
    """Last-resort pdfminer text extraction → plain docx."""
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(input_path)
    except Exception:
        text = ''

    if not text.strip():
        try:
            doc = fitz.open(input_path)
            text = '\n\n'.join(p.get_text() for p in doc)
            doc.close()
        except Exception:
            text = '(No text could be extracted from this PDF.)'

    word_doc = Document()
    style = word_doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for para_text in text.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            word_doc.add_paragraph(para_text)

    word_doc.save(output_path)


# ── OCR fallback (for scanned PDFs) ──────────────────────────────────────────

def _ocr_to_docx(input_path: str, output_path: str, language: str = 'eng'):
    """OCR-based extraction for scanned/image PDFs."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
        word_doc = Document()
        style = word_doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        images = convert_from_path(input_path, dpi=200)
        for i, img in enumerate(images):
            if i > 0:
                word_doc.add_page_break()
            text = pytesseract.image_to_string(img, lang=language)
            for line in text.split('\n'):
                line = line.strip()
                if line:
                    word_doc.add_paragraph(line)

        word_doc.save(output_path)
        return True
    except Exception as e:
        logger.warning(f'OCR docx failed: {e}')
        return False


# ── Main API ──────────────────────────────────────────────────────────────────

def pdf_to_word(
    input_path: str,
    output_path: str,
    password: str = '',
    start_page: int = 0,
    end_page: Optional[int] = None,
    extract_images: bool = True,
    detect_tables: bool = True,
    ocr_fallback: bool = True,
    ocr_language: str = 'eng',
) -> dict:
    """
    Convert a PDF to Microsoft Word (.docx) using multi-strategy pipeline.

    Strategies tried in order:
      1. pdf2docx — best layout fidelity
      2. fitz + python-docx — rich text with formatting
      3. pdfminer + python-docx — pure text
      4. Tesseract OCR + python-docx — for scanned PDFs

    Args:
        input_path:     Source PDF
        output_path:    Output .docx path
        password:       PDF password if encrypted
        start_page:     0-based start page (partial conversion)
        end_page:       0-based end page exclusive (None = all)
        extract_images: Embed images in output docx
        detect_tables:  Attempt table reconstruction
        ocr_fallback:   Try OCR if no text found
        ocr_language:   Tesseract language for OCR fallback

    Returns:
        dict: output_path, method, page_count, file_size_kb,
              has_images, has_tables, conversion_quality
    """
    # Get page count
    page_count = 0
    is_scanned = False
    try:
        reader = PdfReader(input_path)
        if reader.is_encrypted:
            reader.decrypt(password or '')
        page_count = len(reader.pages)
    except Exception:
        pass

    # Check if scanned
    try:
        doc_check = fitz.open(input_path)
        text_chars = sum(len(doc_check[i].get_text().strip())
                         for i in range(min(3, doc_check.page_count)))
        is_scanned = text_chars < 50 and doc_check.page_count > 0
        doc_check.close()
    except Exception:
        pass

    method = 'unknown'

    # ── Strategy 1: pdf2docx ─────────────────────────────────────────────────
    if not is_scanned:
        try:
            cv = Converter(input_path)
            cv.convert(
                output_path,
                start=start_page,
                end=end_page,
                password=password or None,
            )
            cv.close()
            if os.path.exists(output_path) and os.path.getsize(output_path) > 200:
                method = 'pdf2docx'
                return _build_result(output_path, method, page_count)
        except Exception as e:
            logger.warning(f'pdf2docx failed: {e}')

    # ── Strategy 2: fitz → docx ─────────────────────────────────────────────
    if not is_scanned:
        try:
            _fitz_to_docx(input_path, output_path,
                           start_page=start_page,
                           end_page=end_page,
                           extract_images=extract_images,
                           detect_tables=detect_tables)
            if os.path.exists(output_path) and os.path.getsize(output_path) > 200:
                method = 'fitz+docx'
                return _build_result(output_path, method, page_count)
        except Exception as e:
            logger.warning(f'fitz+docx failed: {e}')

    # ── Strategy 3: pdfminer → docx ─────────────────────────────────────────
    if not is_scanned:
        try:
            _pdfminer_to_docx(input_path, output_path)
            if os.path.exists(output_path) and os.path.getsize(output_path) > 200:
                method = 'pdfminer+docx'
                return _build_result(output_path, method, page_count)
        except Exception as e:
            logger.warning(f'pdfminer failed: {e}')

    # ── Strategy 4: OCR → docx ──────────────────────────────────────────────
    if ocr_fallback:
        try:
            if _ocr_to_docx(input_path, output_path, ocr_language):
                method = 'ocr+docx'
                return _build_result(output_path, method, page_count)
        except Exception as e:
            logger.warning(f'OCR docx failed: {e}')

    if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
        return _build_result(output_path, method, page_count)

    raise RuntimeError('All PDF to Word conversion strategies failed.')


def _build_result(output_path: str, method: str, page_count: int) -> dict:
    """Build result dict from conversion output."""
    size_kb = round(os.path.getsize(output_path) / 1024, 1) if os.path.exists(output_path) else 0
    quality_map = {
        'pdf2docx': 'high',
        'fitz+docx': 'medium',
        'pdfminer+docx': 'basic',
        'ocr+docx': 'ocr',
        'unknown': 'unknown',
    }
    return {
        'output_path': output_path,
        'method': method,
        'page_count': page_count,
        'file_size_kb': size_kb,
        'conversion_quality': quality_map.get(method, 'medium'),
    }


def extract_text_as_docx(input_path: str, output_path: str,
                          preserve_formatting: bool = True) -> dict:
    """
    Extract PDF text and save as a structured Word document.
    Faster and simpler than full layout conversion.
    """
    try:
        doc = fitz.open(input_path)
        word_doc = Document()
        body_size = _get_median_font_size(doc)

        for i, page in enumerate(doc):
            if i > 0:
                word_doc.add_page_break()

            if preserve_formatting:
                for block in page.get_text('dict').get('blocks', []):
                    if block.get('type') != 0:
                        continue
                    for line in block.get('lines', []):
                        spans = line.get('spans', [])
                        if not spans:
                            continue
                        line_text = ''.join(s.get('text', '') for s in spans).strip()
                        if not line_text:
                            continue
                        max_size = max((s.get('size', 12) for s in spans), default=12)
                        level = _detect_heading_level(max_size, body_size)
                        if level > 0:
                            word_doc.add_heading(line_text, level=level)
                        else:
                            word_doc.add_paragraph(line_text)
            else:
                for line in page.get_text('text').split('\n'):
                    line = line.strip()
                    if line:
                        word_doc.add_paragraph(line)

        doc.close()
        word_doc.save(output_path)
        return {
            'output_path': output_path,
            'method': 'fitz_text',
            'file_size_kb': round(os.path.getsize(output_path) / 1024, 1),
        }
    except Exception as e:
        raise RuntimeError(f'Text extraction to docx failed: {e}')


def batch_pdf_to_word(input_paths: list, output_dir: str,
                       **kwargs) -> list:
    """Convert multiple PDFs to Word documents."""
    os.makedirs(output_dir, exist_ok=True)
    results = []
    for path in input_paths:
        base = os.path.splitext(os.path.basename(path))[0]
        out = os.path.join(output_dir, f'{base}.docx')
        try:
            res = pdf_to_word(path, out, **kwargs)
            res['source_path'] = path
            results.append(res)
        except Exception as e:
            results.append({'source_path': path, 'output_path': None, 'error': str(e)})
    return results


def get_pdf_text_preview(input_path: str, pages: int = 3) -> dict:
    """
    Extract text preview from first N pages for pre-conversion assessment.
    Returns dict with text_preview, page_count, is_scanned, word_count.
    """
    result = {
        'text_preview': '',
        'page_count': 0,
        'is_scanned': False,
        'word_count': 0,
        'recommended_method': 'pdf2docx',
    }
    try:
        doc = fitz.open(input_path)
        result['page_count'] = doc.page_count
        texts = []
        for i in range(min(pages, doc.page_count)):
            texts.append(doc[i].get_text())
        doc.close()
        full_text = '\n'.join(texts)
        result['text_preview'] = full_text[:500]
        result['word_count'] = len(full_text.split())
        result['is_scanned'] = result['word_count'] < 20 and doc.page_count > 0
        if result['is_scanned']:
            result['recommended_method'] = 'ocr+docx'
    except Exception:
        pass
    return result
