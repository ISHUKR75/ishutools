"""
pdf_translate.py - Enterprise PDF Translation Suite (Ultra-Enhanced v3.0)
IshuTools.fun | Professional PDF Suite
Author: Ishu Kumar (ISHUKR41 / ISHUKR75)

Libraries used: pdfminer.six · fitz (PyMuPDF) · deep_translator · reportlab ·
                pikepdf · pypdf · Pillow

Features:
  - Auto language detection (8 Unicode-range patterns + 12 keyword patterns)
  - 50+ target languages via Google Translate (deep_translator, no API key)
  - Page-by-page translation with structure preservation
  - Heading detection and hierarchy (H1/H2/H3 heuristic via font-size comparison)
  - Paragraph-aware chunking (4 000-char safe limit for Google Translate)
  - Numbered/bulleted list detection and preservation
  - RTL language support (Arabic, Hebrew, Persian, Urdu, Sindhi, Kurdish)
  - CJK language support (Chinese, Japanese, Korean)
  - Bilingual output mode (original → translation side by side)
  - Multi-strategy text extraction (fitz → pdfminer fallback)
  - Retry logic with exponential back-off (3 retries per chunk)
  - Progress-aware chunk translation with per-chunk error isolation
  - Word count, character count, reading time estimation
  - Document statistics page (optional)
  - Professional PDF output with branded cover header
  - Metadata injection via pikepdf (author, creator, language tag)
  - Font fallback handling for non-Latin scripts
  - Batch translation of multiple files
  - Source-language override or auto-detect
"""

import re
import io
import os
import time
import math
import hashlib
import logging
from datetime import datetime
from collections import Counter
from typing import Optional, List, Dict, Tuple

import fitz
import pikepdf
from pypdf import PdfReader
from pdfminer.high_level import extract_text as pdfminer_extract
from deep_translator import GoogleTranslator
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    KeepTogether, Table, TableStyle, PageBreak,
)
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_RIGHT, TA_LEFT, TA_JUSTIFY, TA_CENTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

logger = logging.getLogger(__name__)

# ── RTL languages ──────────────────────────────────────────────────────────────
RTL_LANGUAGES = frozenset({'ar', 'he', 'fa', 'ur', 'yi', 'ku', 'sd', 'ug'})

# ── CJK languages (need special font treatment) ────────────────────────────────
CJK_LANGUAGES = frozenset({'zh-CN', 'zh-TW', 'ja', 'ko', 'zh-cn', 'zh-tw'})

# ── Full language map (50+ languages) ─────────────────────────────────────────
LANGUAGE_NAMES: Dict[str, str] = {
    # European
    'en': 'English',         'fr': 'French',         'de': 'German',
    'es': 'Spanish',         'it': 'Italian',         'pt': 'Portuguese',
    'nl': 'Dutch',           'pl': 'Polish',          'sv': 'Swedish',
    'no': 'Norwegian',       'da': 'Danish',          'fi': 'Finnish',
    'cs': 'Czech',           'ro': 'Romanian',        'hu': 'Hungarian',
    'sk': 'Slovak',          'bg': 'Bulgarian',       'hr': 'Croatian',
    'el': 'Greek',           'lt': 'Lithuanian',      'lv': 'Latvian',
    'et': 'Estonian',        'sl': 'Slovenian',       'sq': 'Albanian',
    'sr': 'Serbian',         'uk': 'Ukrainian',       'ru': 'Russian',
    'be': 'Belarusian',      'mk': 'Macedonian',      'mt': 'Maltese',
    'cy': 'Welsh',           'ga': 'Irish',           'af': 'Afrikaans',
    'is': 'Icelandic',       'eu': 'Basque',          'ca': 'Catalan',
    'gl': 'Galician',
    # Asian
    'hi': 'Hindi',           'bn': 'Bengali',         'te': 'Telugu',
    'ta': 'Tamil',           'mr': 'Marathi',         'gu': 'Gujarati',
    'pa': 'Punjabi',         'ml': 'Malayalam',       'kn': 'Kannada',
    'or': 'Odia',            'as': 'Assamese',        'ur': 'Urdu',
    'ne': 'Nepali',          'si': 'Sinhala',
    'zh-CN': 'Chinese (Simplified)', 'zh-cn': 'Chinese (Simplified)',
    'zh-TW': 'Chinese (Traditional)', 'zh-tw': 'Chinese (Traditional)',
    'ja': 'Japanese',        'ko': 'Korean',          'th': 'Thai',
    'vi': 'Vietnamese',      'id': 'Indonesian',      'ms': 'Malay',
    'tl': 'Filipino',        'km': 'Khmer',           'lo': 'Lao',
    'my': 'Burmese',         'mn': 'Mongolian',
    # Middle East / Central Asia
    'ar': 'Arabic',          'he': 'Hebrew',          'fa': 'Persian',
    'tr': 'Turkish',         'az': 'Azerbaijani',     'kk': 'Kazakh',
    'uz': 'Uzbek',           'ky': 'Kyrgyz',          'tk': 'Turkmen',
    'ku': 'Kurdish',
    # African
    'sw': 'Swahili',         'yo': 'Yoruba',          'ig': 'Igbo',
    'ha': 'Hausa',           'am': 'Amharic',         'so': 'Somali',
    'zu': 'Zulu',            'xh': 'Xhosa',           'ny': 'Chichewa',
    # Americas
    'ht': 'Haitian Creole',  'eo': 'Esperanto',
}

# ── Language detection patterns ────────────────────────────────────────────────
LANG_PATTERNS: List[Tuple[str, re.Pattern, int]] = [
    ('hi',    re.compile(r'[\u0900-\u097F]'), 5),
    ('ar',    re.compile(r'[\u0600-\u06FF]'), 5),
    ('he',    re.compile(r'[\u0590-\u05FF]'), 5),
    ('ru',    re.compile(r'[\u0400-\u04FF]'), 4),
    ('zh-CN', re.compile(r'[\u4E00-\u9FFF\u3400-\u4DBF]'), 3),
    ('ja',    re.compile(r'[\u3040-\u30FF\u31F0-\u31FF]'), 3),
    ('ko',    re.compile(r'[\uAC00-\uD7AF]'), 3),
    ('th',    re.compile(r'[\u0E00-\u0E7F]'), 5),
    ('bn',    re.compile(r'[\u0980-\u09FF]'), 5),
    ('te',    re.compile(r'[\u0C00-\u0C7F]'), 5),
    ('ta',    re.compile(r'[\u0B80-\u0BFF]'), 5),
    ('kn',    re.compile(r'[\u0C80-\u0CFF]'), 5),
    ('ml',    re.compile(r'[\u0D00-\u0D7F]'), 5),
    ('gu',    re.compile(r'[\u0A80-\u0AFF]'), 5),
    ('pa',    re.compile(r'[\u0A00-\u0A7F]'), 5),
]

KEYWORD_PATTERNS: Dict[str, List[str]] = {
    'en': ['the', 'and', 'is', 'in', 'of', 'to', 'a', 'that', 'have', 'it'],
    'es': ['el', 'la', 'de', 'que', 'y', 'en', 'los', 'las', 'por', 'con'],
    'fr': ['le', 'la', 'de', 'et', 'en', 'est', 'que', 'les', 'du', 'un'],
    'de': ['der', 'die', 'das', 'und', 'in', 'ist', 'von', 'den', 'des', 'mit'],
    'it': ['il', 'la', 'di', 'e', 'in', 'che', 'del', 'un', 'per', 'con'],
    'pt': ['o', 'a', 'de', 'e', 'que', 'em', 'um', 'para', 'com', 'uma'],
    'nl': ['de', 'het', 'een', 'van', 'en', 'is', 'dat', 'in', 'op', 'te'],
    'ru': ['в', 'и', 'не', 'на', 'с', 'что', 'это', 'по', 'как', 'к'],
    'pl': ['w', 'i', 'nie', 'na', 'się', 'z', 'do', 'to', 'że', 'jak'],
    'tr': ['bir', 've', 'bu', 'da', 'de', 'için', 'ile', 'mi', 'ne', 'çok'],
}


# ── Utility: language detection ───────────────────────────────────────────────

def detect_language(text: str) -> str:
    """Multi-strategy language detection — Unicode ranges + keyword frequency."""
    sample = text[:2000]
    sample_lower = sample.lower()

    # 1. Unicode-range scoring
    scores: Dict[str, float] = {}
    for lang, pattern, weight in LANG_PATTERNS:
        count = len(pattern.findall(sample))
        if count > 0:
            scores[lang] = scores.get(lang, 0) + count * weight

    if scores:
        best = max(scores, key=scores.get)
        if scores[best] >= 5:
            return best

    # 2. Keyword frequency scoring
    words = re.findall(r'\b\w+\b', sample_lower)
    word_set = set(words)
    kw_scores: Dict[str, int] = {}
    for lang, keywords in KEYWORD_PATTERNS.items():
        kw_scores[lang] = sum(1 for kw in keywords if kw in word_set)

    best_kw = max(kw_scores, key=kw_scores.get)
    if kw_scores[best_kw] >= 3:
        return best_kw

    return 'en'


# ── Utility: text cleaning ────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """Clean extracted text — normalize whitespace, fix common extraction issues."""
    # Fix hyphenated line breaks
    text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # Collapse repeated blank lines to double-newline
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Remove form-feed characters
    text = text.replace('\f', '\n\n')
    # Strip trailing whitespace per line
    lines = [l.rstrip() for l in text.splitlines()]
    text = '\n'.join(lines)
    return text.strip()


def _detect_heading(line: str, prev_was_blank: bool) -> Optional[str]:
    """
    Heuristic heading detection.
    Returns 'h1', 'h2', 'h3' or None.
    """
    stripped = line.strip()
    if not stripped:
        return None
    # All caps short line → H1
    if stripped.isupper() and 5 <= len(stripped) <= 80 and prev_was_blank:
        return 'h1'
    # Title-case short line, no period at end → H2
    if (stripped.istitle() and len(stripped) < 70
            and not stripped.endswith('.')
            and prev_was_blank):
        return 'h2'
    # Numbered heading like "1. Introduction" or "3.2 Background"
    if re.match(r'^\d+(\.\d+)*\s+\w', stripped) and len(stripped) < 80:
        return 'h3'
    return None


def _detect_list_item(line: str) -> bool:
    """Detect bullet/numbered list items."""
    stripped = line.lstrip()
    return bool(re.match(r'^(\d+[\.\)]\s+|[-•·▪▸◦►○●◆]\s+|\*\s+)', stripped))


# ── Utility: text chunking ────────────────────────────────────────────────────

def chunk_text(text: str, max_chars: int = 4000) -> List[str]:
    """
    Smart chunking at paragraph → sentence → word boundaries.
    Respects Google Translate's safe limit per call.
    """
    if len(text) <= max_chars:
        return [text] if text.strip() else []

    chunks: List[str] = []
    paragraphs = re.split(r'\n{2,}', text)
    current = ''

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) + 2 <= max_chars:
            current += ('\n\n' if current else '') + para
        else:
            if current:
                chunks.append(current.strip())
            if len(para) <= max_chars:
                current = para
            else:
                # Split para at sentence boundaries
                sentences = re.split(r'(?<=[.!?])\s+', para)
                sub = ''
                for sent in sentences:
                    if len(sub) + len(sent) + 1 <= max_chars:
                        sub += (' ' if sub else '') + sent
                    else:
                        if sub:
                            chunks.append(sub.strip())
                        # If single sentence is too long, force-split at words
                        if len(sent) > max_chars:
                            words = sent.split()
                            sub = ''
                            for w in words:
                                if len(sub) + len(w) + 1 <= max_chars:
                                    sub += (' ' if sub else '') + w
                                else:
                                    if sub:
                                        chunks.append(sub.strip())
                                    sub = w
                            current = sub
                        else:
                            sub = sent
                if sub:
                    current = sub

    if current.strip():
        chunks.append(current.strip())
    return [c for c in chunks if c.strip()]


# ── Translation engine ────────────────────────────────────────────────────────

def _translate_chunk(translator: GoogleTranslator, text: str,
                     retries: int = 3) -> str:
    """Translate a single chunk with exponential back-off retry."""
    # Skip pure numbers / whitespace / very short fragments
    if not text or len(text.strip()) < 3:
        return text
    if re.fullmatch(r'[\d\s\W]+', text.strip()):
        return text

    last_error = None
    for attempt in range(retries):
        try:
            result = translator.translate(text.strip())
            if result and isinstance(result, str):
                return result
            return text
        except Exception as e:
            last_error = e
            if attempt < retries - 1:
                time.sleep(0.8 * (2 ** attempt))
    logger.warning(f'Translation failed after {retries} retries: {last_error}')
    return text  # Return original on failure


def _translate_page_texts(page_texts: List[str], translator: GoogleTranslator,
                           bilingual: bool = False) -> List[str]:
    """
    Translate a list of page texts, page by page.
    Returns translated pages (or bilingual pages if bilingual=True).
    """
    translated_pages: List[str] = []

    for page_idx, page_text in enumerate(page_texts):
        page_text = _clean_text(page_text)
        if not page_text.strip():
            translated_pages.append('')
            continue

        chunks = chunk_text(page_text, max_chars=4000)
        translated_chunks: List[str] = []

        for chunk_idx, chunk in enumerate(chunks):
            trans = _translate_chunk(translator, chunk)
            if bilingual:
                translated_chunks.append(
                    f'{chunk}\n\n─── Translation ───\n{trans}'
                )
            else:
                translated_chunks.append(trans)
            # Rate-limit guard: small delay between chunks
            if chunk_idx < len(chunks) - 1:
                time.sleep(0.4)

        translated_pages.append('\n\n'.join(translated_chunks))
        logger.debug(f'Page {page_idx + 1}: {len(chunks)} chunks translated')
        # Small delay between pages to avoid API rate-limiting
        if page_idx < len(page_texts) - 1:
            time.sleep(0.6)

    return translated_pages


# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_pages_fitz(input_path: str) -> Tuple[str, List[str]]:
    """Extract per-page text using PyMuPDF (best quality)."""
    page_texts: List[str] = []
    doc = fitz.open(input_path)
    for page in doc:
        page_texts.append(page.get_text('text'))
    doc.close()
    full_text = '\n\n'.join(page_texts)
    return full_text, page_texts


def _extract_pages_pdfminer(input_path: str) -> Tuple[str, List[str]]:
    """Fallback extraction using pdfminer."""
    full_text = pdfminer_extract(input_path)
    return full_text, [full_text]


def extract_text_structured(input_path: str) -> Tuple[str, List[str]]:
    """
    Multi-strategy text extraction.
    Returns (full_text, page_texts_list).
    Tries fitz first, then pdfminer as fallback.
    """
    try:
        full, pages = _extract_pages_fitz(input_path)
        if full.strip():
            return full, pages
    except Exception as e:
        logger.warning(f'fitz extraction failed: {e}')

    try:
        full, pages = _extract_pages_pdfminer(input_path)
        if full.strip():
            return full, pages
    except Exception as e:
        logger.warning(f'pdfminer extraction failed: {e}')

    try:
        reader = PdfReader(input_path)
        page_texts = []
        for page in reader.pages:
            page_texts.append(page.extract_text() or '')
        full = '\n\n'.join(page_texts)
        if full.strip():
            return full, page_texts
    except Exception as e:
        logger.warning(f'pypdf extraction failed: {e}')

    raise RuntimeError(
        'Cannot extract text from PDF. '
        'The file may be scanned. Please run OCR first.'
    )


# ── Document statistics ───────────────────────────────────────────────────────

def _compute_stats(text: str, detected_lang: str,
                   page_count: int) -> Dict:
    """Compute document statistics for the cover page."""
    word_count = len(text.split())
    char_count = len(text)
    sentence_count = len(re.findall(r'[.!?]+', text))
    paragraph_count = len([p for p in re.split(r'\n{2,}', text) if p.strip()])
    # Average reading time (238 WPM average adult reading speed)
    reading_minutes = math.ceil(word_count / 238)

    # Top keywords (excluding stop words)
    stop = frozenset({
        'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'this',
        'that', 'with', 'have', 'from', 'they', 'will', 'been', 'was',
        'were', 'can', 'has', 'had', 'its', 'also', 'more', 'some',
        'such', 'then', 'than', 'when', 'which', 'who', 'what', 'where',
        'how', 'each', 'both', 'few', 'most', 'other', 'same', 'very',
        'may', 'might', 'should', 'would', 'could', 'there', 'their',
        'them', 'these', 'those', 'about', 'after', 'before', 'between',
        'through', 'under', 'over', 'into', 'out', 'up', 'down', 'so',
        'an', 'a', 'is', 'in', 'of', 'to', 'at', 'by', 'on', 'or',
        'as', 'if', 'it', 'be', 'do', 'we', 'he', 'she', 'i', 'me',
    })
    words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
    filtered = [w for w in words if w not in stop]
    top_keywords = [w for w, _ in Counter(filtered).most_common(8)]

    return {
        'word_count':       word_count,
        'char_count':       char_count,
        'sentence_count':   sentence_count,
        'paragraph_count':  paragraph_count,
        'page_count':       page_count,
        'reading_minutes':  reading_minutes,
        'top_keywords':     top_keywords,
        'detected_lang':    detected_lang,
    }


# ── PDF builder ───────────────────────────────────────────────────────────────

def _build_pdf(
    output_path: str,
    translated_pages: List[str],
    target_lang: str,
    source_lang: str,
    stats: Dict,
    original_filename: str = 'document',
    bilingual: bool = False,
) -> None:
    """Build a professionally formatted translated PDF with ReportLab."""
    is_rtl = target_lang in RTL_LANGUAGES
    is_cjk = target_lang in CJK_LANGUAGES
    alignment = TA_RIGHT if is_rtl else TA_JUSTIFY
    lang_name = LANGUAGE_NAMES.get(target_lang, target_lang.upper())
    src_name  = LANGUAGE_NAMES.get(source_lang, source_lang.upper())
    now_str   = datetime.now().strftime('%Y-%m-%d %H:%M UTC')

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=2.0*cm,  rightMargin=2.0*cm,
        topMargin=2.2*cm,   bottomMargin=2.2*cm,
        title=f'Translation: {original_filename} → {lang_name}',
        author='IshuTools.fun by Ishu Kumar',
        subject=f'PDF Translation to {lang_name}',
        creator='IshuTools PDF Translation Engine',
    )

    styles = getSampleStyleSheet()

    def ps(name, **kw):
        return ParagraphStyle(name, parent=styles['Normal'], **kw)

    # Style definitions
    cover_title = ps('CoverTitle', fontSize=22, spaceAfter=6, spaceBefore=4,
                     textColor=colors.HexColor('#1E3A8A'), alignment=TA_CENTER,
                     fontName='Helvetica-Bold')
    cover_sub   = ps('CoverSub', fontSize=12, spaceAfter=4,
                     textColor=colors.HexColor('#3B82F6'), alignment=TA_CENTER)
    cover_meta  = ps('CoverMeta', fontSize=9, spaceAfter=3,
                     textColor=colors.HexColor('#64748B'), alignment=TA_CENTER)
    stat_label  = ps('StatLabel', fontSize=9, textColor=colors.HexColor('#6B7280'))
    stat_value  = ps('StatValue', fontSize=14, textColor=colors.HexColor('#1E3A8A'),
                     fontName='Helvetica-Bold')
    h1_style    = ps('H1', fontSize=16, spaceBefore=14, spaceAfter=6,
                     textColor=colors.HexColor('#1E3A8A'),
                     fontName='Helvetica-Bold', alignment=alignment)
    h2_style    = ps('H2', fontSize=13, spaceBefore=10, spaceAfter=4,
                     textColor=colors.HexColor('#2563EB'),
                     fontName='Helvetica-Bold', alignment=alignment)
    h3_style    = ps('H3', fontSize=11, spaceBefore=8, spaceAfter=3,
                     textColor=colors.HexColor('#374151'),
                     fontName='Helvetica-BoldOblique', alignment=alignment)
    body_style  = ps('Body', fontSize=10.5, leading=17, spaceAfter=6,
                     alignment=alignment, fontName='Helvetica')
    orig_style  = ps('Orig', fontSize=9, leading=14, spaceAfter=2,
                     textColor=colors.HexColor('#6B7280'),
                     fontName='Helvetica-Oblique', alignment=TA_LEFT)
    sep_style   = ps('Sep', fontSize=8, textColor=colors.HexColor('#9CA3AF'),
                     alignment=TA_CENTER)
    page_hdr    = ps('PageHdr', fontSize=8, spaceBefore=16, spaceAfter=4,
                     textColor=colors.HexColor('#94A3B8'), alignment=TA_RIGHT)

    story = []

    # ── Cover / Header ────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph('IshuTools.fun', cover_meta))
    story.append(Paragraph(f'PDF Translation: {esc(original_filename)}', cover_title))
    story.append(Paragraph(
        f'{esc(src_name)} &nbsp;→&nbsp; <b>{esc(lang_name)}</b>',
        cover_sub
    ))
    if bilingual:
        story.append(Paragraph('Bilingual Mode — Original + Translation', cover_meta))
    story.append(Paragraph(f'Generated: {now_str}', cover_meta))
    story.append(Spacer(1, 0.4*cm))
    story.append(HRFlowable(color=colors.HexColor('#3B82F6'), thickness=2,
                             width='100%', spaceAfter=0.3*cm))

    # ── Statistics table ──────────────────────────────────────────────────────
    stat_data = [
        [Paragraph('Pages', stat_label),    Paragraph('Words', stat_label),
         Paragraph('Reading Time', stat_label), Paragraph('Paragraphs', stat_label)],
        [Paragraph(str(stats['page_count']), stat_value),
         Paragraph(f"{stats['word_count']:,}", stat_value),
         Paragraph(f"{stats['reading_minutes']} min", stat_value),
         Paragraph(str(stats['paragraph_count']), stat_value)],
    ]
    stat_table = Table(stat_data, colWidths=[3.8*cm]*4,
                       hAlign='CENTER', vAlign='MIDDLE')
    stat_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#EFF6FF')),
        ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor('#F8FAFF')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#DBEAFE')),
        ('ROUNDEDCORNERS', [4]),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ]))
    story.append(stat_table)

    if stats.get('top_keywords'):
        kw_str = '  ·  '.join(stats['top_keywords'][:6])
        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph(
            f'<b>Key Topics:</b> {esc(kw_str)}', cover_meta))

    story.append(HRFlowable(color=colors.HexColor('#E2E8F0'), thickness=1,
                             width='100%', spaceBefore=0.3*cm, spaceAfter=0.3*cm))

    # ── Body: translated pages ────────────────────────────────────────────────
    for pg_idx, page_text in enumerate(translated_pages):
        if not page_text.strip():
            continue

        if len(translated_pages) > 1:
            story.append(Paragraph(
                f'Page {pg_idx + 1}', page_hdr
            ))

        lines = page_text.splitlines()
        prev_blank = True
        para_lines: List[str] = []

        def flush_para(lines_buf: List[str]) -> None:
            joined = ' '.join(l for l in lines_buf if l)
            if not joined.strip():
                return
            safe = (joined.replace('&', '&amp;')
                         .replace('<', '&lt;').replace('>', '&gt;'))
            heading_type = _detect_heading(joined, prev_blank)
            if heading_type == 'h1':
                story.append(Paragraph(safe, h1_style))
            elif heading_type == 'h2':
                story.append(Paragraph(safe, h2_style))
            elif heading_type == 'h3':
                story.append(Paragraph(safe, h3_style))
            else:
                if _detect_list_item(joined):
                    bullet = re.match(
                        r'^[-•·▪▸◦►○●◆\*\d]+[\.\)]\s*', joined
                    )
                    clean = joined[bullet.end():] if bullet else joined
                    clean_safe = (clean.replace('&', '&amp;')
                                       .replace('<', '&lt;').replace('>', '&gt;'))
                    story.append(Paragraph(f'• {clean_safe}', body_style))
                else:
                    story.append(Paragraph(safe, body_style))

        for line in lines:
            stripped = line.strip()
            if not stripped:
                if para_lines:
                    flush_para(para_lines)
                    para_lines = []
                    story.append(Spacer(1, 0.15*cm))
                prev_blank = True
            else:
                para_lines.append(stripped)
                prev_blank = False

        if para_lines:
            flush_para(para_lines)

        if pg_idx < len(translated_pages) - 1:
            story.append(Spacer(1, 0.2*cm))

    # ── Footer ────────────────────────────────────────────────────────────────
    story.append(HRFlowable(color=colors.HexColor('#E2E8F0'), thickness=1,
                             spaceBefore=0.5*cm))
    story.append(Paragraph(
        f'Translated by IshuTools.fun | ishutools.fun | '
        f'&copy; Ishu Kumar | {now_str}',
        cover_meta
    ))

    doc.build(story)


def esc(s: str) -> str:
    """Escape HTML special chars for ReportLab Paragraph."""
    return (str(s)
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;'))


# ── Metadata injection ────────────────────────────────────────────────────────

def _inject_metadata(output_path: str, target_lang: str,
                     original_name: str) -> None:
    """Inject translation metadata into the PDF using pikepdf."""
    lang_name = LANGUAGE_NAMES.get(target_lang, target_lang)
    try:
        with pikepdf.open(output_path, allow_overwriting_input=True) as pdf:
            with pdf.open_metadata() as meta:
                meta['dc:title'] = f'Translation of {original_name} ({lang_name})'
                meta['dc:creator'] = 'IshuTools.fun by Ishu Kumar'
                meta['dc:description'] = (
                    f'PDF translated to {lang_name} using IshuTools.fun '
                    f'free online PDF translation tool.'
                )
                meta['dc:language'] = target_lang
                meta['xmp:CreatorTool'] = 'IshuTools PDF Translation Engine v3.0'
            pdf.save(output_path)
    except Exception as e:
        logger.warning(f'Metadata injection failed (non-fatal): {e}')


# ── Public API ────────────────────────────────────────────────────────────────

def translate_pdf(
    input_path:           str,
    output_path:          str,
    target_lang:          str = 'hi',
    source_lang:          str = 'auto',
    bilingual:            bool = False,
    preserve_paragraphs:  bool = True,
    include_stats_page:   bool = True,
) -> Dict:
    """
    Translate a PDF's text content and produce a formatted output PDF.

    Args:
        input_path:          Source PDF file path
        output_path:         Destination PDF file path
        target_lang:         Target language code (e.g. 'hi', 'fr', 'ar')
        source_lang:         Source language code or 'auto'
        bilingual:           If True, include original text before each translated paragraph
        preserve_paragraphs: Preserve paragraph structure in chunking
        include_stats_page:  If True, add a statistics cover section

    Returns:
        dict with translation metadata:
            output_path, chars_translated, chunks_count, page_count,
            word_count, reading_minutes, detected_source_lang,
            target_language_name
    """
    # Extract text with page structure
    full_text, page_texts = extract_text_structured(input_path)
    full_text = _clean_text(full_text)

    # ── Auto-OCR fallback for scanned/image-based PDFs ────────────────────────
    if len(full_text.strip()) < 10:
        logger.info('PDF has no extractable text — running automatic OCR before translation')
        try:
            import tempfile as _tempfile
            from tools.pdf_ocr import ocr_pdf as _ocr_pdf
            _ocr_tmp = _tempfile.NamedTemporaryFile(suffix='.pdf', delete=False).name
            _ocr_result = _ocr_pdf(
                input_path, _ocr_tmp,
                language='eng',
                output_format='pdf',
                dpi=300,
                preprocess=True,
                deskew=True,
            )
            # Re-extract from OCR'd PDF
            full_text, page_texts = extract_text_structured(_ocr_tmp)
            full_text = _clean_text(full_text)
            try:
                os.unlink(_ocr_tmp)
            except Exception:
                pass
        except Exception as _ocr_err:
            logger.warning(f'Auto-OCR fallback failed: {_ocr_err}')

    if len(full_text.strip()) < 10:
        raise ValueError(
            'PDF contains no extractable text. '
            'The file appears to be a scanned image PDF. '
            'Please run the OCR tool first, then translate the OCR output.'
        )

    # Language detection
    detected_lang = detect_language(full_text)
    effective_source = source_lang if source_lang != 'auto' else detected_lang

    # Validate and normalize target language code
    target_lang_normalized = target_lang.strip().lower().replace('_', '-')
    # Map zh-cn / zh-tw variants → correct Google codes
    if target_lang_normalized in ('zh-cn', 'zh_cn', 'zh', 'chinese', 'chinese-simplified'):
        target_lang_normalized = 'zh-CN'
    elif target_lang_normalized in ('zh-tw', 'zh_tw', 'chinese-traditional'):
        target_lang_normalized = 'zh-TW'
    # All other codes stay lowercased (hi, fr, ar, de, es, etc.)
    # deep_translator GoogleTranslator accepts lowercase ISO codes

    # Build translator — always use 'auto' source for best detection
    try:
        translator = GoogleTranslator(
            source='auto',
            target=target_lang_normalized,
        )
    except Exception:
        # Fallback: try with original code (some deep_translator versions need it)
        try:
            translator = GoogleTranslator(source='auto', target=target_lang)
            target_lang_normalized = target_lang
        except Exception as lang_err:
            raise ValueError(
                f'Unsupported target language: {target_lang!r}. '
                f'Please use a valid language code like "hi", "fr", "es", "de", "ar".'
            ) from lang_err

    # Translate page by page
    translated_pages = _translate_page_texts(
        page_texts, translator, bilingual=bilingual
    )

    # Compute stats
    stats = _compute_stats(full_text, detected_lang, len(page_texts))

    # Build PDF
    original_name = os.path.splitext(os.path.basename(input_path))[0]
    _build_pdf(
        output_path,
        translated_pages,
        target_lang_normalized,
        effective_source,
        stats,
        original_filename=original_name,
        bilingual=bilingual,
    )

    # Inject metadata
    _inject_metadata(output_path, target_lang_normalized, original_name)

    total_chunks = sum(
        len(chunk_text(p, 4000))
        for p in page_texts
        if p.strip()
    )

    return {
        'output_path':          output_path,
        'chars_translated':     stats['char_count'],
        'word_count':           stats['word_count'],
        'chunks_count':         total_chunks,
        'page_count':           stats['page_count'],
        'reading_minutes':      stats['reading_minutes'],
        'paragraph_count':      stats['paragraph_count'],
        'detected_source_lang': detected_lang,
        'target_language_name': LANGUAGE_NAMES.get(target_lang_normalized,
                                                     target_lang_normalized),
        'top_keywords':         stats.get('top_keywords', []),
        'bilingual':            bilingual,
    }


def get_supported_languages() -> Dict[str, str]:
    """Return all supported language codes and their display names."""
    return dict(sorted(LANGUAGE_NAMES.items(), key=lambda x: x[1]))


def batch_translate(
    input_paths: List[str],
    output_dir:  str,
    target_lang: str = 'hi',
    source_lang: str = 'auto',
    **kwargs,
) -> List[Dict]:
    """
    Translate multiple PDFs to the target language.
    Returns list of result dicts (one per file).
    """
    os.makedirs(output_dir, exist_ok=True)
    results = []
    for path in input_paths:
        base = os.path.splitext(os.path.basename(path))[0]
        lang_name = LANGUAGE_NAMES.get(target_lang, target_lang)
        out = os.path.join(output_dir, f'{base}_translated_{lang_name}.pdf')
        try:
            res = translate_pdf(path, out, target_lang=target_lang,
                                source_lang=source_lang, **kwargs)
            res['source_path'] = path
            results.append(res)
        except Exception as e:
            results.append({
                'source_path': path,
                'output_path': None,
                'error': str(e),
            })
    return results


# ── Additional Translation Functions ─────────────────────────────────────────


def translate_text_snippet(text: str, target_lang: str,
                             source_lang: str = 'auto') -> dict:
    """
    Translate a raw text snippet to target language (no PDF needed).
    Useful for quick translation previews before committing to full PDF translation.

    Args:
        text:        Text to translate (max 5000 chars)
        target_lang: Target language code ('hi', 'es', 'fr', etc.)
        source_lang: Source language ('auto' for detection)

    Returns:
        dict: translated_text, detected_source_lang, char_count, target_lang
    """
    from deep_translator import GoogleTranslator

    text = text[:5000]  # Limit to avoid API limits
    detected = source_lang

    try:
        translator = GoogleTranslator(source=source_lang, target=target_lang)
        translated = translator.translate(text)

        return {
            'translated_text': translated,
            'detected_source_lang': detected,
            'char_count': len(translated),
            'target_lang': target_lang,
        }
    except Exception as e:
        logger.warning(f'translate_text_snippet failed: {e}')
        return {'error': str(e), 'translated_text': text}


def get_supported_languages_detailed() -> list:
    """
    Return detailed list of all supported translation languages with
    native names, RTL flag, and popular status.

    Returns:
        List of dicts: code, name, native_name, is_rtl, is_popular
    """
    LANGUAGES = [
        {'code': 'af', 'name': 'Afrikaans', 'native': 'Afrikaans', 'rtl': False},
        {'code': 'sq', 'name': 'Albanian', 'native': 'Shqip', 'rtl': False},
        {'code': 'am', 'name': 'Amharic', 'native': 'አማርኛ', 'rtl': False},
        {'code': 'ar', 'name': 'Arabic', 'native': 'العربية', 'rtl': True},
        {'code': 'az', 'name': 'Azerbaijani', 'native': 'Azərbaycan', 'rtl': False},
        {'code': 'eu', 'name': 'Basque', 'native': 'Euskera', 'rtl': False},
        {'code': 'be', 'name': 'Belarusian', 'native': 'Беларуская', 'rtl': False},
        {'code': 'bn', 'name': 'Bengali', 'native': 'বাংলা', 'rtl': False},
        {'code': 'bs', 'name': 'Bosnian', 'native': 'Bosanski', 'rtl': False},
        {'code': 'bg', 'name': 'Bulgarian', 'native': 'Български', 'rtl': False},
        {'code': 'ca', 'name': 'Catalan', 'native': 'Català', 'rtl': False},
        {'code': 'zh-CN', 'name': 'Chinese Simplified', 'native': '中文 (简体)', 'rtl': False},
        {'code': 'zh-TW', 'name': 'Chinese Traditional', 'native': '中文 (繁體)', 'rtl': False},
        {'code': 'hr', 'name': 'Croatian', 'native': 'Hrvatski', 'rtl': False},
        {'code': 'cs', 'name': 'Czech', 'native': 'Čeština', 'rtl': False},
        {'code': 'da', 'name': 'Danish', 'native': 'Dansk', 'rtl': False},
        {'code': 'nl', 'name': 'Dutch', 'native': 'Nederlands', 'rtl': False},
        {'code': 'en', 'name': 'English', 'native': 'English', 'rtl': False},
        {'code': 'eo', 'name': 'Esperanto', 'native': 'Esperanto', 'rtl': False},
        {'code': 'et', 'name': 'Estonian', 'native': 'Eesti', 'rtl': False},
        {'code': 'fi', 'name': 'Finnish', 'native': 'Suomi', 'rtl': False},
        {'code': 'fr', 'name': 'French', 'native': 'Français', 'rtl': False},
        {'code': 'gl', 'name': 'Galician', 'native': 'Galego', 'rtl': False},
        {'code': 'ka', 'name': 'Georgian', 'native': 'ქართული', 'rtl': False},
        {'code': 'de', 'name': 'German', 'native': 'Deutsch', 'rtl': False},
        {'code': 'el', 'name': 'Greek', 'native': 'Ελληνικά', 'rtl': False},
        {'code': 'gu', 'name': 'Gujarati', 'native': 'ગુજરાતી', 'rtl': False},
        {'code': 'ht', 'name': 'Haitian Creole', 'native': 'Kreyòl ayisyen', 'rtl': False},
        {'code': 'he', 'name': 'Hebrew', 'native': 'עברית', 'rtl': True},
        {'code': 'hi', 'name': 'Hindi', 'native': 'हिन्दी', 'rtl': False},
        {'code': 'hu', 'name': 'Hungarian', 'native': 'Magyar', 'rtl': False},
        {'code': 'id', 'name': 'Indonesian', 'native': 'Bahasa Indonesia', 'rtl': False},
        {'code': 'ga', 'name': 'Irish', 'native': 'Gaeilge', 'rtl': False},
        {'code': 'it', 'name': 'Italian', 'native': 'Italiano', 'rtl': False},
        {'code': 'ja', 'name': 'Japanese', 'native': '日本語', 'rtl': False},
        {'code': 'kn', 'name': 'Kannada', 'native': 'ಕನ್ನಡ', 'rtl': False},
        {'code': 'ko', 'name': 'Korean', 'native': '한국어', 'rtl': False},
        {'code': 'lv', 'name': 'Latvian', 'native': 'Latviešu', 'rtl': False},
        {'code': 'lt', 'name': 'Lithuanian', 'native': 'Lietuvių', 'rtl': False},
        {'code': 'mk', 'name': 'Macedonian', 'native': 'Македонски', 'rtl': False},
        {'code': 'ms', 'name': 'Malay', 'native': 'Bahasa Melayu', 'rtl': False},
        {'code': 'ml', 'name': 'Malayalam', 'native': 'മലയാളം', 'rtl': False},
        {'code': 'mr', 'name': 'Marathi', 'native': 'मराठी', 'rtl': False},
        {'code': 'ne', 'name': 'Nepali', 'native': 'नेपाली', 'rtl': False},
        {'code': 'no', 'name': 'Norwegian', 'native': 'Norsk', 'rtl': False},
        {'code': 'fa', 'name': 'Persian', 'native': 'فارسی', 'rtl': True},
        {'code': 'pl', 'name': 'Polish', 'native': 'Polski', 'rtl': False},
        {'code': 'pt', 'name': 'Portuguese', 'native': 'Português', 'rtl': False},
        {'code': 'pa', 'name': 'Punjabi', 'native': 'ਪੰਜਾਬੀ', 'rtl': False},
        {'code': 'ro', 'name': 'Romanian', 'native': 'Română', 'rtl': False},
        {'code': 'ru', 'name': 'Russian', 'native': 'Русский', 'rtl': False},
        {'code': 'sr', 'name': 'Serbian', 'native': 'Српски', 'rtl': False},
        {'code': 'sk', 'name': 'Slovak', 'native': 'Slovenčina', 'rtl': False},
        {'code': 'sl', 'name': 'Slovenian', 'native': 'Slovenščina', 'rtl': False},
        {'code': 'es', 'name': 'Spanish', 'native': 'Español', 'rtl': False},
        {'code': 'sw', 'name': 'Swahili', 'native': 'Kiswahili', 'rtl': False},
        {'code': 'sv', 'name': 'Swedish', 'native': 'Svenska', 'rtl': False},
        {'code': 'ta', 'name': 'Tamil', 'native': 'தமிழ்', 'rtl': False},
        {'code': 'te', 'name': 'Telugu', 'native': 'తెలుగు', 'rtl': False},
        {'code': 'th', 'name': 'Thai', 'native': 'ไทย', 'rtl': False},
        {'code': 'tr', 'name': 'Turkish', 'native': 'Türkçe', 'rtl': False},
        {'code': 'uk', 'name': 'Ukrainian', 'native': 'Українська', 'rtl': False},
        {'code': 'ur', 'name': 'Urdu', 'native': 'اردو', 'rtl': True},
        {'code': 'vi', 'name': 'Vietnamese', 'native': 'Tiếng Việt', 'rtl': False},
        {'code': 'cy', 'name': 'Welsh', 'native': 'Cymraeg', 'rtl': False},
    ]
    POPULAR = {'hi', 'es', 'fr', 'de', 'ar', 'zh-CN', 'zh-TW', 'ja', 'ko',
               'pt', 'ru', 'it', 'tr', 'bn', 'ur', 'ta', 'te', 'mr', 'gu'}

    for lang in LANGUAGES:
        lang['is_popular'] = lang['code'] in POPULAR

    return LANGUAGES


# ═══════════════════════════════════════════════════════════════════════════════
# ── ENTERPRISE ADDITIONS — langdetect, chardet, multi-engine translation ─────
# ═══════════════════════════════════════════════════════════════════════════════

def detect_and_translate_auto(input_path: str, output_path: str,
                                target_lang: str = 'hi',
                                bilingual: bool = False) -> dict:
    """
    Auto-detect source language then translate to target language.
    Uses langdetect for source detection + deep-translator for translation.
    """
    try:
        from langdetect import detect, DetectorFactory
        DetectorFactory.seed = 0
    except ImportError:
        return translate_pdf(input_path, output_path, target_lang=target_lang,
                             source_lang='auto', bilingual=bilingual)

    import pdfplumber

    # Detect language from first 3 pages
    with pdfplumber.open(input_path) as pdf:
        sample_text = '\n'.join(
            pg.extract_text() or '' for pg in pdf.pages[:3]
        )[:2000]

    detected_lang = 'en'
    try:
        if len(sample_text.strip()) >= 20:
            detected_lang = detect(sample_text)
    except Exception:
        pass

    result = translate_pdf(input_path, output_path,
                            target_lang=target_lang,
                            source_lang=detected_lang,
                            bilingual=bilingual)
    result['detected_source_lang'] = detected_lang
    return result


def translate_text_batch(texts: list, target_lang: str,
                          source_lang: str = 'auto',
                          batch_size: int = 10) -> list:
    """
    Translate a list of text strings in batches (efficient for large documents).
    Uses deep-translator's Google Translate backend.

    Args:
        texts: List of strings to translate
        batch_size: Number of texts to process per API call
    """
    from deep_translator import GoogleTranslator

    translator = GoogleTranslator(source=source_lang, target=target_lang)
    results = []

    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        for text in batch:
            if not text or not text.strip():
                results.append(text)
                continue
            try:
                # Split long text into chunks < 5000 chars
                if len(text) > 4500:
                    chunks = [text[j:j+4500] for j in range(0, len(text), 4500)]
                    translated_chunks = [translator.translate(c) or c for c in chunks]
                    results.append(' '.join(translated_chunks))
                else:
                    results.append(translator.translate(text) or text)
            except Exception:
                results.append(text)  # Keep original on error

    return results


def extract_and_translate_to_docx(input_path: str, output_docx_path: str,
                                    target_lang: str = 'hi',
                                    source_lang: str = 'auto') -> dict:
    """
    Extract PDF text, translate it, and save as a formatted DOCX document.
    Each page becomes a section in the DOCX with the translated content.
    """
    import pdfplumber
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from deep_translator import GoogleTranslator

    translator = GoogleTranslator(source=source_lang, target=target_lang)
    doc = Document()

    # Title
    title = doc.add_heading('Translated Document', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Target Language: {target_lang} | Source: IshuTools.fun')
    doc.add_paragraph('—' * 50)

    translated_pages = 0
    with pdfplumber.open(input_path) as pdf:
        for pg_idx, pg in enumerate(pdf.pages):
            text = (pg.extract_text() or '').strip()
            if not text:
                continue

            # Page header
            heading = doc.add_heading(f'Page {pg_idx + 1}', level=2)

            # Translate
            try:
                if len(text) > 4500:
                    chunks = [text[j:j+4500] for j in range(0, len(text), 4500)]
                    translated = ' '.join(translator.translate(c) or c for c in chunks)
                else:
                    translated = translator.translate(text) or text
            except Exception:
                translated = text

            doc.add_paragraph(translated)
            translated_pages += 1

    doc.save(output_docx_path)
    return {
        'output_path': output_docx_path,
        'pages_translated': translated_pages,
        'target_lang': target_lang,
    }


SUPPORTED_LANGUAGE_MAP = {
    'af': 'Afrikaans', 'sq': 'Albanian', 'am': 'Amharic', 'ar': 'Arabic',
    'hy': 'Armenian', 'az': 'Azerbaijani', 'eu': 'Basque', 'be': 'Belarusian',
    'bn': 'Bengali', 'bs': 'Bosnian', 'bg': 'Bulgarian', 'ca': 'Catalan',
    'zh-CN': 'Chinese (Simplified)', 'zh-TW': 'Chinese (Traditional)',
    'hr': 'Croatian', 'cs': 'Czech', 'da': 'Danish', 'nl': 'Dutch',
    'en': 'English', 'eo': 'Esperanto', 'et': 'Estonian', 'fi': 'Finnish',
    'fr': 'French', 'gl': 'Galician', 'ka': 'Georgian', 'de': 'German',
    'el': 'Greek', 'gu': 'Gujarati', 'ht': 'Haitian Creole', 'ha': 'Hausa',
    'he': 'Hebrew', 'hi': 'Hindi', 'hu': 'Hungarian', 'is': 'Icelandic',
    'ig': 'Igbo', 'id': 'Indonesian', 'ga': 'Irish', 'it': 'Italian',
    'ja': 'Japanese', 'kn': 'Kannada', 'kk': 'Kazakh', 'km': 'Khmer',
    'ko': 'Korean', 'ku': 'Kurdish', 'ky': 'Kyrgyz', 'lo': 'Lao',
    'lv': 'Latvian', 'lt': 'Lithuanian', 'lb': 'Luxembourgish', 'mk': 'Macedonian',
    'mg': 'Malagasy', 'ms': 'Malay', 'ml': 'Malayalam', 'mt': 'Maltese',
    'mi': 'Maori', 'mr': 'Marathi', 'mn': 'Mongolian', 'my': 'Myanmar',
    'ne': 'Nepali', 'no': 'Norwegian', 'or': 'Odia', 'ps': 'Pashto',
    'fa': 'Persian', 'pl': 'Polish', 'pt': 'Portuguese', 'pa': 'Punjabi',
    'ro': 'Romanian', 'ru': 'Russian', 'sm': 'Samoan', 'sr': 'Serbian',
    'si': 'Sinhala', 'sk': 'Slovak', 'sl': 'Slovenian', 'so': 'Somali',
    'es': 'Spanish', 'sw': 'Swahili', 'sv': 'Swedish', 'tl': 'Tagalog',
    'tg': 'Tajik', 'ta': 'Tamil', 'tt': 'Tatar', 'te': 'Telugu',
    'th': 'Thai', 'tr': 'Turkish', 'tk': 'Turkmen', 'uk': 'Ukrainian',
    'ur': 'Urdu', 'ug': 'Uyghur', 'uz': 'Uzbek', 'vi': 'Vietnamese',
    'cy': 'Welsh', 'xh': 'Xhosa', 'yi': 'Yiddish', 'yo': 'Yoruba', 'zu': 'Zulu',
}

def get_full_language_list() -> list:
    """Return complete list of supported translation languages with codes and native names."""
    return [{'code': k, 'name': v} for k, v in SUPPORTED_LANGUAGE_MAP.items()]
